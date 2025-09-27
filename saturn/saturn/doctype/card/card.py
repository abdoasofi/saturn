# -*- coding: utf-8 -*-
"""
card.py

Complete server-side module for the `Card` DocType.
Place this file at: saturn/saturn/doctype/card/card.py

Features:
- Card.before_save to collect filterable fields
- Whitelisted helpers: set_card_filters, save_filters
- DOCX export: export_labels_docx and export_labels_docx_for_names
  - Renders `card.layout` as Jinja per document using frappe.render_template
  - Inserts images found in layout (only local /files/ images)
  - Generates QR codes using qrcode library
  - Builds a simple docx using python-docx

Requirements (install inside bench env):
    pip install python-docx qrcode pillow

Notes:
- The exporter copies generated file to public/files and returns a url for the client.
- For advanced HTML -> DOCX transformation you may want to use a dedicated library.
"""

import frappe
from frappe.model.document import Document
import json
import io
import os
import uuid
import shutil
import re
from html import unescape

# External libs
from docx import Document as DocxDocument
from docx.shared import Mm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import qrcode

from frappe.utils import get_url, get_site_path


class Card(Document):
    def before_save(self):
        """
        Collect usable fields from the target DocType (doctype_card) to present
        as filter setters in the MultiSelectDialog on the client-side.
        """
        if not self.doctype_card:
            return

        try:
            fields = frappe.get_all(
                "DocField",
                fields=["fieldname", "default"],
                filters={"parent": self.doctype_card},
                or_filters={"in_standard_filter": 1, "in_list_view": 1},
            )
            fields = {f["fieldname"]: f.get("default") for f in fields}
            self.fields = json.dumps(fields)
        except Exception as e:
            # don't block save if something odd happens, but log for debugging
            frappe.log_error(f"Card.before_save failed: {e}", "Card.before_save")


@frappe.whitelist()
def set_card_filters(card_name, filters, starting_value_for_autonumbering=None):
    """
    Save the user's selection (filters) to a small Card Use Log for current user.
    Used by the client JS to remember what the user selected before printing/export.
    """
    user = frappe.session.user
    doctype = "Card Use Log"
    card_log_name = frappe.db.exists(doctype, {"card": card_name, "user": user})

    if card_log_name:
        card_log = frappe.get_doc(doctype, card_log_name)
    else:
        card_log = frappe.get_doc({"doctype": doctype, "card": card_name, "user": user})

    card_log.filters = filters
    if starting_value_for_autonumbering:
        card_log.starting_value_for_autonumbering = starting_value_for_autonumbering

    card_log.save(ignore_permissions=True)
    frappe.db.commit()


@frappe.whitelist()
def save_filters(doc_name, filters):
    """Simple server helper to persist filters on the Card doc itself."""
    doc = frappe.get_doc("Card", doc_name)
    doc.filters = filters
    doc.save(ignore_permissions=True)
    return True


# ------------------ Helpers for HTML -> docx minimal conversion ------------------

def _html_to_paragraphs(html):
    """
    Clean HTML then convert to plain paragraphs.
    - remove <style>...</style>, <script>...</script>, and HTML comments
    - normalize <br>, </p>, </div>, <h*> to newlines
    - strip remaining tags and unescape entities
    - return list of non-empty lines
    """
    if not html:
        return []

    # remove style blocks (CSS), script blocks, and HTML comments
    # (?is) => DOTALL + IGNORECASE
    html = re.sub(r'(?is)<\s*style[^>]*>.*?<\s*/\s*style\s*>', '', html)
    html = re.sub(r'(?is)<\s*script[^>]*>.*?<\s*/\s*script\s*>', '', html)
    html = re.sub(r'(?is)<!--.*?-->', '', html)

    # Normalize known break tags to newlines
    html = re.sub(r'(?i)<\s*br\s*/?\s*>', '\n', html)
    html = re.sub(r'(?i)</\s*p\s*>', '\n', html)
    html = re.sub(r'(?i)</\s*div\s*>', '\n', html)
    html = re.sub(r'(?i)<\s*(h[1-6])[^>]*>', '\n', html)

    # Remove all remaining tags
    text = re.sub(r'<[^>]+>', '', html)
    text = unescape(text)

    # Split and sanitize
    lines = [ln.strip() for ln in text.splitlines()]
    return [ln for ln in lines if ln]


def _extract_img_srcs(html):
    """Return list of src values found in <img ...> tags in the HTML."""
    if not html:
        return []
    srcs = re.findall(r'<img[^>]+src=["\']([^"\']+)["\']', html, flags=re.IGNORECASE)
    if not srcs:
        srcs = re.findall(r'<img[^>]+src=([^>\s]+)', html, flags=re.IGNORECASE)
    return srcs


# ------------------ DOCX Exporters ------------------

@frappe.whitelist()
def export_labels_docx(card_name):
    """
    Export DOCX using the filters saved in Card Use Log for the current user.
    Returns: {"url": "/files/..."}
    """
    card = frappe.get_doc("Card", card_name)

    user = frappe.session.user
    card_log_name = frappe.db.exists("Card Use Log", {"card": card_name, "user": user})
    if not card_log_name:
        frappe.throw("No Card Use Log found for the current user. Please select items first using Print Cards dialog.")

    card_log = frappe.get_doc("Card Use Log", card_log_name)
    raw_filters = card_log.filters or card.filters or ""
    try:
        filters_obj = json.loads(raw_filters) if raw_filters else {}
    except Exception:
        try:
            filters_obj = json.loads(raw_filters.replace("'", '"'))
        except Exception:
            frappe.throw("Unable to parse filters from Card Use Log.")

    doctype = card.doctype_card
    filters = {}
    if isinstance(filters_obj, dict):
        filters = filters_obj
    elif isinstance(filters_obj, list):
        filters = {"name": ["in", filters_obj]}
    else:
        frappe.throw("Unsupported filters format.")

    docs = frappe.get_all(doctype, filters=filters, fields=["*"], limit_page_length=2000)
    if not docs:
        frappe.throw("No documents found for selected filters.")

    # Build docx
    docx = DocxDocument()
    section = docx.sections[0]
    section.top_margin = Mm(5)
    section.bottom_margin = Mm(5)
    section.left_margin = Mm(5)
    section.right_margin = Mm(5)

    def g(d, *keys):
        for k in keys:
            if d.get(k):
                return d.get(k)
        return ""

    for d in docs:
        # create layout table (2 columns)
        table = docx.add_table(rows=1, cols=2)
        table.autofit = False
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        left_cell.width = Mm(120)
        right_cell.width = Mm(40)

        # LEFT side: use minimal HTML rendering if card.layout exists; else fallback fields
        if card.layout:
            try:
                rendered = frappe.render_template(card.layout, {"doc": d, "doc_card": d})
            except Exception as e:
                frappe.log_error(f"Template render failed for {d.get('name')}: {e}", "export_labels_docx")
                rendered = ""

            paragraphs = _html_to_paragraphs(rendered)
            if paragraphs:
                p = left_cell.paragraphs[0]
                # first line as title
                p.add_run(paragraphs[0]).bold = True
                p.runs[0].font.size = Pt(12)
                for l in paragraphs[1:]:
                    p2 = left_cell.add_paragraph()
                    p2.add_run(l).font.size = Pt(9)
            else:
                title = g(d, "item_name", "name", "title")
                p = left_cell.paragraphs[0]
                p.add_run(str(title)).bold = True
                p.runs[0].font.size = Pt(12)
        else:
            title = g(d, "item_name", "name", "title")
            p = left_cell.paragraphs[0]
            p.add_run(str(title)).bold = True
            p.runs[0].font.size = Pt(12)

        # RIGHT side: try to insert first local image found in rendered HTML, then QR
        image_added = False
        try:
            # prefer image fields commonly used
            image_field = None
            for f in ("image", "item_image", "photo"):
                if d.get(f):
                    image_field = d.get(f)
                    break

            if image_field:
                img_path = None
                if image_field.startswith('/files/') or image_field.startswith('files/'):
                    fname = os.path.basename(image_field)
                    candidate = get_site_path('public', 'files', fname)
                    if os.path.exists(candidate):
                        right_cell.paragraphs[0].add_run().add_picture(candidate, width=Mm(30))
                        image_added = True

            if not image_added and card.layout:
                # try to find <img> in rendered HTML
                rendered = frappe.render_template(card.layout, {"doc": d, "doc_card": d})
                srcs = _extract_img_srcs(rendered)
                for src in srcs:
                    if src.startswith('/files/') or src.startswith('files/'):
                        fname = os.path.basename(src)
                        candidate = get_site_path('public', 'files', fname)
                        if os.path.exists(candidate):
                            right_cell.paragraphs[0].add_run().add_picture(candidate, width=Mm(30))
                            image_added = True
                            break
        except Exception as e:
            frappe.log_error(f"Image insertion failed for {d.get('name')}: {e}", "export_labels_docx")

        # QR generation
        qr_value = d.get('name') or d.get('item_code') or ''
        try:
            qr = qrcode.QRCode(box_size=3, border=1)
            qr.add_data(str(qr_value))
            qr.make(fit=True)
            img = qr.make_image(fill_color='black', back_color='white')
            bio = io.BytesIO()
            img.save(bio, format='PNG')
            bio.seek(0)
            right_cell.add_paragraph().add_run().add_picture(bio, width=Mm(30))
        except Exception:
            frappe.log_error(f"QR generation failed for {qr_value}", "export_labels_docx")

        # code under QR
        code_par = right_cell.add_paragraph()
        code_par.add_run(str(qr_value)).font.size = Pt(8)
        code_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # spacer
        docx.add_paragraph()

    # Save result
    filename = f"labels_{card_name}_{uuid.uuid4().hex[:8]}.docx"
    tmp_path = os.path.join('/tmp', filename)
    docx.save(tmp_path)
    public_dest = get_site_path('public', 'files', filename)
    try:
        shutil.copy(tmp_path, public_dest)
    except Exception as e:
        frappe.log_error(f"Failed copying docx to public/files: {e}", "export_labels_docx")
        frappe.throw('Failed to save the generated file on server. Check permissions.')

    file_url = get_url('/files/' + filename)
    return {"url": file_url}


@frappe.whitelist()
def export_labels_docx_for_names(card_name, names):
    """
    Export DOCX for specific document names with fixed product image size:
    - product image: width=3inches, height=2inches
    """
    import io as _io
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.shared import Mm, Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    try:
        names_list = json.loads(names) if isinstance(names, str) else names
    except Exception:
        frappe.throw("Invalid 'names' parameter. Expecting JSON array.")

    if not names_list:
        frappe.throw("No items selected for export.")

    card = frappe.get_doc('Card', card_name)
    doctype = card.doctype_card
    if not doctype:
        frappe.throw('Card has no target DocType set (doctype_card).')

    docs = frappe.get_all(doctype, filters={"name": ["in", names_list]}, fields=["*"], limit_page_length=2000)
    if not docs:
        frappe.throw('No documents found for the given names.')

    docx = DocxDocument()
    section = docx.sections[0]
    section.top_margin = Mm(5)
    section.bottom_margin = Mm(5)
    section.left_margin = Mm(5)
    section.right_margin = Mm(5)

    for d in docs:
        table = docx.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # try to set column widths (best-effort)
        try:
            table.columns[0].width = Mm(120)
            table.columns[1].width = Mm(40)
        except Exception:
            pass

        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # LEFT: product image with fixed size 3in x 2in
        image_field = None
        for f in ('image', 'item_image', 'photo'):
            if d.get(f):
                image_field = d.get(f)
                break

        if image_field:
            try:
                if image_field.startswith('/files/') or image_field.startswith('files/'):
                    fname = os.path.basename(image_field)
                    candidate = get_site_path('public', 'files', fname)
                    if os.path.exists(candidate):
                        # here we set fixed width and height in Inches
                        left_cell.paragraphs[0].add_run().add_picture(candidate, width=Inches(4.3), height=Inches(1.2))
                    else:
                        left_cell.paragraphs[0].add_run("No image file")
                elif image_field.startswith('http'):
                    # remote image: optionally download or leave placeholder
                    left_cell.paragraphs[0].add_run("Image (remote)")
                else:
                    candidate = get_site_path('public', 'files', os.path.basename(image_field))
                    if os.path.exists(candidate):
                        left_cell.paragraphs[0].add_run().add_picture(candidate, width=Inches(4.3), height=Inches(1.2))
                    else:
                        left_cell.paragraphs[0].add_run("No image")
            except Exception:
                left_cell.paragraphs[0].add_run("No image")
        else:
            # keep empty cell or add placeholder blank paragraph to keep layout
            left_cell.paragraphs[0].add_run("")

        # RIGHT: QR (we keep as before)
        qr_value =  d.get('item_code') or ''
        try:
            qr = qrcode.QRCode(box_size=6, border=1)
            qr.add_data(str(qr_value))
            qr.make(fit=True)
            img = qr.make_image(fill_color='black', back_color='white')
            bio = _io.BytesIO()
            img.save(bio, format='PNG')
            bio.seek(0)
            # you can adjust QR size if you want; here we use 30mm ~ 1.18in
            right_cell.paragraphs[0].add_run().add_picture(bio, width=Mm(30))
        except Exception:
            right_cell.paragraphs[0].add_run("")

        small_par = right_cell.add_paragraph()
        small_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        small_run = small_par.add_run(str(d.get('saturn_code') or ''))
        small_run.bold = True
        small_run.font.size = Pt(10)

        # merged product name row
        row_name = table.rows[1]
        cell_name = row_name.cells[0].merge(row_name.cells[1])
        p_name = cell_name.paragraphs[0]
        p_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_name = p_name.add_run(str(d.get('item_name') or d.get('name') or ''))
        run_name.bold = True
        run_name.font.size = Pt(10)

        # merged footer row
        row_footer = table.rows[2]
        cell_footer = row_footer.cells[0].merge(row_footer.cells[1])
        p_footer = cell_footer.paragraphs[0]
        p_footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_text = "SC SATURN ACCESORY SLTEL MAGAZIN 0724888429"
        f_run = p_footer.add_run(footer_text)
        f_run.font.size = Pt(8)

        docx.add_paragraph()

    filename = f"labels_{card_name}_{uuid.uuid4().hex[:8]}.docx"
    tmp_path = os.path.join('/tmp', filename)
    docx.save(tmp_path)
    public_dest = get_site_path('public', 'files', filename)
    try:
        shutil.copy(tmp_path, public_dest)
    except Exception as e:
        frappe.log_error(f"Failed copying docx to public/files: {e}", "export_labels_docx_for_names")
        frappe.throw('Failed to save the generated file on server. Check permissions.')

    file_url = get_url('/files/' + filename)
    return {"url": file_url}